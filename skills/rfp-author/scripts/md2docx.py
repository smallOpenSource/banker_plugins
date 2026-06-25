#!/usr/bin/env python3
# md -> docx : headings / tables / bullets / blockquote / bold, Korean font, fixed-width tables, page numbers
# usage: python md2docx.py <input.md> <output.docx> [font]
#   - 이 서버엔 pandoc/soffice 부재 → python-docx 사용 (pip install python-docx)
#   - 표는 고정 레이아웃 + tblGrid/gridCol(twips)로 비균등폭(맨왼쪽 fit-content). cell.width만으로는 균등폭으로 렌더됨.
import sys
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

if len(sys.argv) < 3:
    print("usage: python md2docx.py <input.md> <output.docx> [font]")
    sys.exit(2)
SRC = sys.argv[1]
OUT = sys.argv[2]
FONT = sys.argv[3] if len(sys.argv) > 3 else "Malgun Gothic"  # 한글 폰트(서버에 설치된 것으로 교체 가능)

def set_korean(style):
    style.font.name = FONT
    style.font.size = Pt(10)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts'); rpr.append(rfonts)
    for a in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):  # eastAsia 필수(누락 시 한글 깨짐)
        rfonts.set(qn(a), FONT)

def add_runs(p, text):
    for part in re.split(r'(\*\*[^*]+\*\*)', text):
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            p.add_run(part[2:-2].replace('`', '')).bold = True
        else:
            p.add_run(part.replace('`', ''))

def is_sep(cells):
    nonempty = [c.strip() for c in cells if c.strip()]
    return bool(nonempty) and all(re.fullmatch(r':?-{1,}:?', c) for c in nonempty)

def set_table_widths(t):
    # 균등 너비 금지: 맨 왼쪽 컬럼은 내용에 맞춰 최소폭(개행 없이), 나머지 컬럼이 잔여 폭 분배
    if not t.rows:
        return
    ncol = len(t.columns)
    if ncol < 2:
        return
    TOTAL = 6.5  # 본문 폭(inch)
    col0 = 0.6
    for row in t.rows:
        w = 0.22  # 셀 패딩
        for ch in row.cells[0].text:
            w += 0.14 if ord(ch) > 0x2000 else 0.08  # CJK ~0.14, ascii ~0.08
        col0 = max(col0, w)
    col0 = min(col0, 3.2)  # 과도 폭 방지(데이터 컬럼 확보)
    rest = (TOTAL - col0) / (ncol - 1)
    t.autofit = False
    t.allow_autofit = False
    tblPr = t._tbl.tblPr
    lay = tblPr.find(qn('w:tblLayout'))
    if lay is None:
        lay = OxmlElement('w:tblLayout'); tblPr.append(lay)
    lay.set(qn('w:type'), 'fixed')  # 고정 레이아웃(지정 폭 준수)
    widths = [col0] + [rest] * (ncol - 1)
    grid = t._tbl.find(qn('w:tblGrid'))
    if grid is not None:
        for i, gc in enumerate(grid.findall(qn('w:gridCol'))):
            if i < len(widths):
                gc.set(qn('w:w'), str(int(widths[i] * 1440)))  # inch→twips, 고정레이아웃 실폭(★ 이게 핵심)
    for i, col in enumerate(t.columns):
        for cell in col.cells:
            cell.width = Inches(widths[i])

doc = Document()
set_korean(doc.styles['Normal'])

lines = open(SRC, encoding='utf-8').read().split('\n')
i, n = 0, len(lines)
while i < n:
    s = lines[i].strip()
    if not s:
        i += 1; continue
    if s.startswith('|'):
        block = []
        while i < n and lines[i].strip().startswith('|'):
            block.append(lines[i].strip()); i += 1
        def split_row(tl):
            return [c.strip() for c in tl.strip('|').split('|')]
        rows = [split_row(tl) for tl in block]
        rows = [r for r in rows if not is_sep(r)]
        if rows:
            ncol = max(len(r) for r in rows)
            t = doc.add_table(rows=0, cols=ncol); t.style = 'Table Grid'
            for ri, r in enumerate(rows):
                rc = t.add_row().cells
                for ci in range(ncol):
                    para = rc[ci].paragraphs[0]
                    add_runs(para, r[ci] if ci < len(r) else '')
                    if ri == 0:
                        for run in para.runs:
                            run.bold = True
            set_table_widths(t)
            doc.add_paragraph()  # spacer after table (표와 다음 단락 겹침 방지)
        continue
    if s.startswith('### '):
        doc.add_heading(s[4:], level=2)
    elif s.startswith('## '):
        doc.add_heading(s[3:], level=1)
    elif s.startswith('# '):
        h0 = doc.add_heading('', level=0); add_runs(h0, s[2:])
        for r in h0.runs:
            r.font.size = Pt(20)
    elif s.startswith('> '):
        p = doc.add_paragraph(); add_runs(p, s[2:])
        for run in p.runs:
            run.italic = True
            run.font.size = Pt(8)
    elif s.startswith('- '):
        add_runs(doc.add_paragraph(style='List Bullet'), s[2:])
    else:
        add_runs(doc.add_paragraph(), s)
    i += 1

# centered page number in footer
fp = doc.sections[0].footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = fp.add_run()
for tag, kw in (('w:fldChar', {'w:fldCharType': 'begin'}), (None, None), ('w:fldChar', {'w:fldCharType': 'end'})):
    if tag is None:
        el = OxmlElement('w:instrText'); el.set(qn('xml:space'), 'preserve'); el.text = 'PAGE'
    else:
        el = OxmlElement(tag)
        for k, v in kw.items():
            el.set(qn(k), v)
    run._r.append(el)

doc.save(OUT)
print("saved", OUT)
